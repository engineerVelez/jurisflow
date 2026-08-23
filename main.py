from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, Form
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
import shutil
from drive import generar_hash, obtener_o_crear_carpeta, archivo_existe
from drive_auth import get_service
from docx_utils import leer_docx, convertir_doc_a_docx, convertir_pdf_a_docx
from ia import analizar_con_ia, calcular_confianza, chatear_con_ia, PROMPT_VERSION
from fastapi import Form
from fastapi import Body
from fastapi.responses import FileResponse
from docx import Document
import os
import sys
import re
import datetime
import base64

# 🔧 FIX: la consola de Windows (cp1252) no puede imprimir emojis/acentos
# raros y UN CRASH de print() rompía el análisis completo (la IA devolvía
# {}). Se fuerza la salida a UTF-8 con reemplazo, para que los print de
# depuración nunca tiren la aplicación.
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

app = FastAPI()

app.mount("/static", StaticFiles(directory="static"), name="static")

# 🔐 LOGIN
security = HTTPBasic()

def verificar(credentials: HTTPBasicCredentials = Depends(security)):
    if credentials.username != "ursula" or credentials.password != "1121":
        raise HTTPException(status_code=401)
    return credentials.username

# 🏠 HOME
@app.get("/", response_class=HTMLResponse)
def home(user: str = Depends(verificar)):
    with open("templates/index.html", "r", encoding="utf-8") as f:
        return f.read()

# 🔧 FIX: un análisis guardado se considera COMPLETO solo si trae la
# estructura actual de la IA (actor/demandado como diccionarios y los
# campos planos de los 5 testigos). Si falta algo, se re-analiza.
def datos_incompletos(datos):
    if not isinstance(datos, dict):
        return True
    # el actor/demandado puede venir como dict (análisis de la IA:
    # {"nombre": ...}) o como texto plano (guardado manual del cliente:
    # "NOMBRE APELLIDO"). Ambos son válidos para reutilizar el archivo.
    if not isinstance(datos.get("actor"), (dict, str)) or not isinstance(datos.get("demandado"), (dict, str)):
        return True
    for i in range(1, 6):
        if datos.get(f"nombre_testigo{i}") is None:
            return True
    return False


# 🔧 FIX: los guardados manuales del usuario (formato plano, con
# "prompt_usado": "") son su fuente de verdad: se cargan SIEMPRE que el
# archivo sea el mismo, sin importar el prompt configurado ni la versión.
# Solo los análisis automáticos de la IA exigen prompt/versión iguales.
def es_guardado_manual(datos_guardados):
    if not isinstance(datos_guardados, dict):
        return False
    if datos_guardados.get("guardado_manual"):
        return True
    datos = datos_guardados.get("datos", {})
    # formato manual: actor es TEXTO plano; formato IA: actor es DICT
    return isinstance(datos.get("actor"), str) and datos_guardados.get("prompt_usado") == ""


# 🔧 FIX: al cargar un guardado manual, el "texto" guardado es el texto
# plano del editor (con los cambios destructivos que el usuario hizo).
# Si el usuario re-subió el archivo ORIGINAL, ese texto ya no coincide
# exactamente con el docx, así que comparamos de forma NORMALIZADA
# (mayúsculas, espacios y saltos) para no re-analizar y perder datos.
def normalizar_para_restaurar(t):
    if not isinstance(t, str):
        return ""
    return re.sub(r"\s+", " ", t).strip().lower()


# 📂 UPLOAD
@app.post("/upload")
def upload(
    file: UploadFile = File(...),
    prompt: str = Form("")
):
    print("🔥 BACKEND FUNCIONA")  # 👈 AGREGA ESTO

    ruta = f"temp_{file.filename}"

    with open(ruta, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # 🔧 NUEVO: soporte .doc (formato antiguo) convirtiendo con Word al subir
    nombre_min = file.filename.lower()
    if nombre_min.endswith(".doc") and not nombre_min.endswith(".docx"):
        ruta = convertir_doc_a_docx(ruta)
        if not ruta:
            raise HTTPException(
                status_code=400,
                detail="No se pudo leer el archivo .doc (Word no pudo convertirlo)."
            )

    # 🔧 NUEVO: soporte .pdf convirtiendo a .docx con Word al subir
    if nombre_min.endswith(".pdf"):
        ruta = convertir_pdf_a_docx(ruta)
        if not ruta:
            raise HTTPException(
                status_code=400,
                detail="No se pudo leer el archivo PDF (Word no pudo convertirlo)."
            )

    texto = leer_docx(ruta)

    # 🔧 FIX: el prompt del usuario (Configurar IA) debe SIEMPRE aplicarse.
    # Si el archivo ya estaba analizado con OTRO prompt o con una versión
    # vieja del prompt del servidor, la caché de Drive se ignora y se
    # vuelve a analizar.
    prompt_usuario = (prompt or "").strip()

    # 🔥 1. GENERAR HASH
    hash_doc = generar_hash(file.filename)

    # 🔥 2. CONECTAR DRIVE
    service = get_service()
    CARPETA_PRINCIPAL = "1m_yHvLo0XKpBavojbHTOw0z2nduzjT7k"

    # 🔥 3. CREAR O BUSCAR
    carpeta_id = obtener_o_crear_carpeta(service, hash_doc, CARPETA_PRINCIPAL)

    # 🔥 4. VER SI YA EXISTE
    if archivo_existe(service, "datos.json", carpeta_id):

        from drive import leer_json

        datos_guardados = leer_json(service, carpeta_id, "datos.json")

        # 🔥 VALIDAR SI EL JSON SIRVE
        if datos_guardados and datos_guardados.get("texto"):

            datos_guardados_datos = datos_guardados.get("datos", {})

            es_manual = es_guardado_manual(datos_guardados)

            if es_manual:
                # 🔧 FIX: el guardado MANUAL es la fuente de verdad del
                # usuario (corrigió campos y se sobrescribió en Drive, sin
                # copias). Se carga SIEMPRE que el archivo sea el mismo
                # (texto normalizado coincide), SIN importar el prompt
                # configurado ni la versión del prompt del servidor. Antes
                # se exigía prompt_usado == prompt_usuario y, si el usuario
                # tenía un prompt configurado, el guardado manual NUNCA se
                # cargaba: la IA re-analizaba y perdía las correcciones.
                texto_referencia = datos_guardados.get("texto_original") or datos_guardados.get("texto", "")

                if normalizar_para_restaurar(texto_referencia) == normalizar_para_restaurar(texto):

                    print("✔ JSON MANUAL → CARGAR")

                    porcentaje, mensaje = calcular_confianza(datos_guardados_datos)

                    return {
                        "mensaje": "Cargado desde Drive",
                        "documento_id": hash_doc,
                        "datos": datos_guardados_datos,
                        "texto": datos_guardados.get("texto", ""),
                        "texto_html": datos_guardados.get("texto_html", ""),
                        "texto_original": texto_referencia,
                        "porcentaje": porcentaje,
                        "mensaje_confianza": mensaje
                    }

                print("⚠️ JSON MANUAL DE OTRO DOCUMENTO → RE-ANALIZAR")

            elif not datos_incompletos(datos_guardados_datos):

                # 🔧 FIX: la caché SOLO sirve si se generó con el MISMO
                # prompt del usuario, la MISMA versión del prompt del
                # servidor Y el MISMO contenido del archivo. La carpeta de
                # Drive se crea con el hash del NOMBRE del archivo: dos
                # documentos con el mismo nombre pero distinto contenido
                # comparten carpeta. Si el texto guardado no coincide con
                # el del archivo actual, la caché es de OTRO documento y
                # se vuelve a analizar (si no, el Entry mostraría valores
                # del documento anterior, no del actual).
                cache_vigente = (
                    datos_guardados.get("texto") == texto
                    and datos_guardados.get("prompt_usado", "") == prompt_usuario
                    and datos_guardados.get("prompt_version") == PROMPT_VERSION
                )

                if cache_vigente:

                    print("✔ YA EXISTE → CARGAR")

                    porcentaje, mensaje = calcular_confianza(datos_guardados_datos)

                    return {
                        "mensaje": "Cargado desde Drive",
                        "documento_id": hash_doc,
                        "datos": datos_guardados_datos,
                        "texto": datos_guardados.get("texto", ""),
                        "texto_html": datos_guardados.get("texto_html", ""),
                        "texto_original": datos_guardados.get("texto", ""),
                        "porcentaje": porcentaje,
                        "mensaje_confianza": mensaje
                    }

                print("⚠️ PROMPT/VERSIÓN CAMBIÓ → RE-ANALIZAR CON IA")

            else:
                print("⚠️ JSON INCOMPLETO → RE-ANALIZAR CON IA")

        else:
            print("⚠️ JSON VACÍO → USAR IA")

    # ===============================
    # 🟢 USAR IA (NUEVO O JSON INCOMPLETO)
    # ===============================

    print("📁 NUEVO O INCOMPLETO → USAR IA")

    datos = analizar_con_ia(texto, prompt_usuario)

    from drive import subir_o_actualizar_json

    subir_o_actualizar_json(service, carpeta_id, "datos.json", {
        "texto": texto,
        "datos": datos,
        "prompt_usado": prompt_usuario,
        "prompt_version": PROMPT_VERSION
    })

    porcentaje, mensaje_confianza = calcular_confianza(datos)

    if len(texto) > 120000:
        mensaje_confianza += " El documento es largo: la IA solo revisó la primera parte del texto."

    return {
        "mensaje": "Archivo procesado",
        "documento_id": hash_doc,
        "datos": datos,
        "texto": texto,
        "texto_original": texto,
        "porcentaje": porcentaje,
        "mensaje_confianza": mensaje_confianza
    }
@app.post("/chat-ia")
def chat_ia(data: dict = Body(...)):

    mensaje = data.get("mensaje", "")
    historial = data.get("historial", [])
    instrucciones = data.get("instrucciones", "")

    if not mensaje or not mensaje.strip():
        raise HTTPException(status_code=400, detail="Mensaje vacío")

    if not isinstance(historial, list):
        historial = []

    # 🔧 FIX: el chat también recibe las instrucciones adicionales que el
    # usuario escribió en Configurar IA, para que la IA las respete al
    # responder y al proponer correcciones.
    respuesta = chatear_con_ia(mensaje.strip(), historial, instrucciones)

    return {"respuesta": respuesta}


@app.post("/chat-ia-stream")
async def chat_ia_stream(data: dict = Body(...)):

    from fastapi.responses import StreamingResponse
    from ia import chatear_con_ia_streaming

    mensaje = data.get("mensaje", "")
    historial = data.get("historial", [])
    instrucciones = data.get("instrucciones", "")
    contexto_documento = data.get("contexto_documento", None)

    if not mensaje or not mensaje.strip():
        raise HTTPException(status_code=400, detail="Mensaje vacío")

    if not isinstance(historial, list):
        historial = []

    def event_generator():
        for chunk in chatear_con_ia_streaming(mensaje.strip(), historial, instrucciones, contexto_documento):
            yield f"data: {chunk}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}
    )


@app.post("/api/ia/crear-documento")
async def crear_documento_ia_endpoint(data: dict = Body(...)):
    from ia import generar_documento_ia
    from docx import Document as DocxDocument
    import io

    instruccion = data.get("instruccion", "")
    historial = data.get("historial", [])

    if not instruccion.strip():
        raise HTTPException(status_code=400, detail="Instrucción vacía")

    resultado = generar_documento_ia(instruccion, historial)

    if not resultado.get("ok"):
        raise HTTPException(status_code=500, detail=resultado.get("error", "Error generando documento"))

    texto = resultado.get("texto", "")
    if not texto:
        raise HTTPException(status_code=500, detail="La IA no generó texto")

    docx_doc = DocxDocument()
    for linea in texto.split("\n"):
        docx_doc.add_paragraph(linea)

    buf = io.BytesIO()
    docx_doc.save(buf)
    file_bytes = buf.getvalue()

    titulo = resultado.get("titulo", "Documento Jurídico")
    nombre_limpio = re.sub(r'[<>:"/\\|?*]', '', titulo).strip()
    if not nombre_limpio:
        nombre_limpio = "documento_generado"
    filename = nombre_limpio + ".docx"

    return {
        "ok": True,
        "texto": texto,
        "titulo": titulo,
        "categoria": resultado.get("categoria", "OTROS"),
        "subcategoria": resultado.get("subcategoria", ""),
        "procedimiento": resultado.get("procedimiento", "NO APLICA"),
        "entries_pendientes": resultado.get("entries_pendientes", []),
        "docx_base64": base64.b64encode(file_bytes).decode("utf-8"),
        "filename": filename
    }


@app.post("/reanalizar")
def reanalizar(data: dict = Body(...)):

    # 🔧 NUEVO: re-analiza el documento ACTUAL (ya cargado en el editor)
    # con el prompt nuevo, para que las correcciones del chat se apliquen
    # de inmediato sin volver a subir el archivo.

    texto = data.get("texto", "")
    prompt = data.get("prompt", "")
    documento_id = data.get("documento_id", "")

    if not texto or not texto.strip():
        raise HTTPException(status_code=400, detail="No hay texto para re-analizar")

    prompt_usuario = (prompt or "").strip()

    service = get_service()
    CARPETA_PRINCIPAL = "1m_yHvLo0XKpBavojbHTOw0z2nduzjT7k"

    # La carpeta se identifica con el MISMO documento_id (hash) del upload,
    # así el datos.json actual se sobrescribe con los datos corregidos.
    hash_doc = documento_id or ""
    carpeta_id = obtener_o_crear_carpeta(service, hash_doc, CARPETA_PRINCIPAL)

    from drive import subir_o_actualizar_json

    datos = analizar_con_ia(texto, prompt_usuario)

    subir_o_actualizar_json(service, carpeta_id, "datos.json", {
        "texto": texto,
        "datos": datos,
        "prompt_usado": prompt_usuario,
        "prompt_version": PROMPT_VERSION
    })

    porcentaje, mensaje_confianza = calcular_confianza(datos)

    if len(texto) > 120000:
        mensaje_confianza += " El documento es largo: la IA solo revisó la primera parte del texto."

    return {
        "mensaje": "Documento reanalizado",
        "documento_id": hash_doc,
        "datos": datos,
        "texto": texto,
        "texto_original": texto,
        "porcentaje": porcentaje,
        "mensaje_confianza": mensaje_confianza
    }


@app.post("/exportar-docx")
def exportar_docx(data: dict = Body(...)):

    texto = data.get("texto", "")
    texto = texto.replace("\\r\\n", "\n").replace("\\n", "\n").replace("\r\n", "\n").replace("\r", "\n")
    texto = texto.replace("  ", " ")

    doc = Document()

    for linea in texto.split("\n"):
        doc.add_paragraph(linea)

    ruta = "documento_generado.docx"
    doc.save(ruta)

    return FileResponse(
        ruta,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="documento_editado.docx"
    )

@app.post("/guardar-html")
def guardar_html(data: dict = Body(...)):

    texto = data.get("texto", "")                    # HTML con resaltados (copia local)
    texto_plano = data.get("texto_plano", texto)     # texto limpio para Drive/reúso
    nombre = data.get("nombre", "documento")
    datos = data.get("datos", {})
    documento_id = data.get("documento_id", "")

    # 🔥 GUARDAR LOCAL (opcional)
    ruta = f"guardados/{nombre}.html"
    os.makedirs("guardados", exist_ok=True)

    with open(ruta, "w", encoding="utf-8") as f:
        f.write(texto)

    # 🔥 GUARDAR EN DRIVE
    from drive import obtener_o_crear_carpeta, subir_o_actualizar_json

    service = get_service()

    # 🔧 FIX: el hash del documento lo envía el frontend (documento_id),
    # el mismo que usó /upload para crear la carpeta en Drive. Antes se
    # usaba "file.filename" que NO existe aquí (daba NameError y el
    # guardado manual en Drive NUNCA funcionaba).
    hash_doc = documento_id or generar_hash(nombre)
    CARPETA_PRINCIPAL = "1m_yHvLo0XKpBavojbHTOw0z2nduzjT7k"

    carpeta_id = obtener_o_crear_carpeta(service, hash_doc, CARPETA_PRINCIPAL)

    subir_o_actualizar_json(service, carpeta_id, "datos.json", {
        "texto": texto_plano,
        "texto_html": texto,
        "texto_original": data.get("texto_original", texto_plano),
        "datos": datos,
        "prompt_usado": "",
        "prompt_version": PROMPT_VERSION,
        "guardado_manual": True
    })

    print("🔥 GUARDADO EN DRIVE")

    return {"ok": True}

@app.get("/cargar-html/{nombre}")
def cargar_html(nombre: str):
    ruta = f"guardados/{nombre}.html"

    if os.path.exists(ruta):
        with open(ruta, "r", encoding="utf-8") as f:
            return {"html": f.read()}

    return {"html": None}


# ============================================================
# 📚 REPOSITORIO DE DOCUMENTOS BASE
# ============================================================

@app.get("/api/documentos-base")
def listar_documentos_base():
    from drive import (
        leer_registro_documentos_base,
        obtener_carpeta_archivos_base,
        leer_datos_documento_base,
    )
    service = get_service()
    registro = leer_registro_documentos_base(service)
    documentos = registro.get("documentos", [])
    for doc in documentos:
        datos_cache = leer_datos_documento_base(service, doc["id"])
        doc["tiene_datos"] = datos_cache is not None and datos_cache.get("texto")
    return {"documentos": documentos}


@app.post("/api/documentos-base/subir")
async def subir_documento_base_endpoint(
    file: UploadFile = File(...),
    nombre: str = Form(""),
    categoria: str = Form(""),
    subcategoria: str = Form(""),
    procedimiento: str = Form(""),
    descripcion: str = Form(""),
):
    from drive import (
        leer_registro_documentos_base,
        guardar_registro_documentos_base,
        subir_documento_base,
        guardar_datos_documento_base,
        guardar_config_documento_base,
    )
    from ia import detectar_variables

    if not file.filename:
        raise HTTPException(status_code=400, detail="No se proporcionó archivo")

    file_bytes = await file.read()

    doc_id = generar_hash(nombre.lower().strip() or file.filename)

    service = get_service()
    subir_documento_base(service, doc_id, file_bytes, file.filename)

    texto = ""
    ruta_temp = f"temp_base_{file.filename}"
    with open(ruta_temp, "wb") as buffer:
        buffer.write(file_bytes)

    nombre_min = file.filename.lower()
    if nombre_min.endswith(".docx"):
        texto = leer_docx(ruta_temp)
    elif nombre_min.endswith(".doc"):
        ruta_conv = convertir_doc_a_docx(ruta_temp)
        if ruta_conv:
            texto = leer_docx(ruta_conv)
    elif nombre_min.endswith(".pdf"):
        ruta_conv = convertir_pdf_a_docx(ruta_temp)
        if ruta_conv:
            texto = leer_docx(ruta_conv)

    if os.path.exists(ruta_temp):
        os.remove(ruta_temp)

    config = None
    if texto:
        try:
            variables = detectar_variables(texto)
            existing_config = leer_config_documento_base(service, doc_id) or {}
            config = dict(existing_config)
            config["variables"] = variables if variables else []
            config["config_version"] = max(config.get("config_version", 1), 2)
            config["fecha_configuracion"] = datetime.datetime.now().isoformat()
            guardar_config_documento_base(service, doc_id, config)
            print(f"🔑 Detectadas {len(config['variables'])} variables en {doc_id}")
            for v in config['variables']:
                print(f"   - {v['key']}: {v['marcador']}")
        except Exception as e:
            print(f"⚠️ Error detectando variables: {e}")
            config = {"variables": [], "config_version": 1}
            guardar_config_documento_base(service, doc_id, config)

        guardar_datos_documento_base(service, doc_id, {
            "texto": texto,
            "texto_original": texto,
            "prompt_usado": "",
            "prompt_version": PROMPT_VERSION,
        })

    registro = leer_registro_documentos_base(service)

    nuevo_doc = {
        "id": doc_id,
        "nombre": nombre.strip() or file.filename,
        "categoria": (categoria or "OTROS").upper().strip(),
        "subcategoria": subcategoria.strip() or None,
        "procedimiento": (procedimiento or "").strip() or None,
        "descripcion": descripcion.strip(),
        "archivo_nombre": file.filename,
        "fecha_creacion": datetime.datetime.now().isoformat(),
        "fecha_modificacion": datetime.datetime.now().isoformat(),
        "version": 1,
    }

    registro["documentos"].append(nuevo_doc)
    _agregarATaxonomia(registro, (categoria or "").upper().strip(), subcategoria.strip() or None, (procedimiento or "").strip() or None)
    guardar_registro_documentos_base(service, registro)

    print(f"📚 DOCUMENTO BASE SUBIDO: {nuevo_doc['nombre']}")

    respuesta = {"ok": True, "documento": nuevo_doc}

    if config:
        respuesta["config"] = config

    if texto:
        from ia import sugerir_clasificacion
        try:
            sugerencia = sugerir_clasificacion(texto)
            if sugerencia:
                respuesta["sugerencia"] = sugerencia
        except Exception as e:
            print(f"⚠️ Error sugiriendo clasificación: {e}")

    return respuesta


@app.delete("/api/documentos-base/{doc_id}")
def eliminar_documento_base_endpoint(doc_id: str):
    from drive import (
        leer_registro_documentos_base,
        guardar_registro_documentos_base,
        eliminar_documento_base_completo,
    )

    service = get_service()
    registro = leer_registro_documentos_base(service)
    documentos = registro.get("documentos", [])
    encontrado = False

    for doc in documentos:
        if doc["id"] == doc_id:
            encontrado = True
            break

    if not encontrado:
        raise HTTPException(status_code=404, detail="Documento base no encontrado")

    try:
        eliminar_documento_base_completo(service, doc_id)
    except Exception as e:
        print(f"⚠️ Error eliminando carpeta en Drive (continuando): {e}")

    registro["documentos"] = [d for d in documentos if d["id"] != doc_id]

    seed_id = generar_hash("demanda_servidumbre_paso_jurisflow")
    if doc_id == seed_id:
        seeds_eliminados = registro.get("seeds_eliminados", [])
        if doc_id not in seeds_eliminados:
            seeds_eliminados.append(doc_id)
            registro["seeds_eliminados"] = seeds_eliminados

    guardar_registro_documentos_base(service, registro)

    print(f"🗑 DOCUMENTO BASE ELIMINADO: {doc_id}")
    return {"ok": True}


@app.patch("/api/documentos-base/{doc_id}")
async def actualizar_documento_base_endpoint(
    doc_id: str,
    nombre: str = Form(""),
    categoria: str = Form(""),
    subcategoria: str = Form(""),
    procedimiento: str = Form(""),
    descripcion: str = Form(""),
    archivo: UploadFile = File(None),
):
    from drive import (
        leer_registro_documentos_base,
        guardar_registro_documentos_base,
        subir_documento_base,
        guardar_datos_documento_base,
        guardar_config_documento_base,
        leer_config_documento_base,
    )
    from ia import detectar_variables

    service = get_service()
    registro = leer_registro_documentos_base(service)
    documentos = registro.get("documentos", [])

    encontrado = False
    for doc in documentos:
        if doc["id"] == doc_id:
            encontrado = True
            if categoria:
                doc["categoria"] = categoria
            if subcategoria is not None:
                doc["subcategoria"] = subcategoria or None
            if procedimiento is not None:
                doc["procedimiento"] = procedimiento or None
            if nombre:
                doc["nombre"] = nombre
            if descripcion is not None:
                doc["descripcion"] = descripcion
            doc["fecha_modificacion"] = datetime.datetime.now().isoformat()
            break

    if not encontrado:
        raise HTTPException(status_code=404, detail="Documento base no encontrado")

    if archivo and archivo.filename:
        file_bytes = await archivo.read()

        subir_documento_base(service, doc_id, file_bytes, archivo.filename)

        texto = ""
        ruta_temp = f"temp_edit_{archivo.filename}"
        with open(ruta_temp, "wb") as buffer:
            buffer.write(file_bytes)

        nombre_min = archivo.filename.lower()
        if nombre_min.endswith(".docx"):
            texto = leer_docx(ruta_temp)
        elif nombre_min.endswith(".doc"):
            ruta_conv = convertir_doc_a_docx(ruta_temp)
            if ruta_conv:
                texto = leer_docx(ruta_conv)
        elif nombre_min.endswith(".pdf"):
            ruta_conv = convertir_pdf_a_docx(ruta_temp)
            if ruta_conv:
                texto = leer_docx(ruta_conv)

        if os.path.exists(ruta_temp):
            os.remove(ruta_temp)

        if texto:
            config = None
            try:
                variables = detectar_variables(texto)
                existing_config = leer_config_documento_base(service, doc_id) or {}
                config = dict(existing_config)
                config["variables"] = variables if variables else []
                config["config_version"] = max(config.get("config_version", 1), 2)
                config["fecha_configuracion"] = datetime.datetime.now().isoformat()
                guardar_config_documento_base(service, doc_id, config)
                print(f"🔑 Detectadas {len(config['variables'])} variables en {doc_id}")
            except Exception as e:
                print(f"⚠️ Error detectando variables: {e}")
                existing_config = leer_config_documento_base(service, doc_id) or {}
                config = dict(existing_config)
                config["variables"] = []
                config["config_version"] = max(config.get("config_version", 1), 2)
                guardar_config_documento_base(service, doc_id, config)

            guardar_datos_documento_base(service, doc_id, {
                "texto": texto,
                "texto_original": texto,
                "prompt_usado": "",
                "prompt_version": PROMPT_VERSION,
            })

            for doc in documentos:
                if doc["id"] == doc_id:
                    doc["archivo_nombre"] = archivo.filename
                    break

        print(f"📄 ARCHIVO REEMPLAZADO en {doc_id}: {archivo.filename}")

    _agregarATaxonomia(registro, categoria.upper().strip() if categoria else "", subcategoria.strip() if subcategoria else None, procedimiento.strip() if procedimiento else None)
    guardar_registro_documentos_base(service, registro)
    print(f"✏️ DOCUMENTO BASE ACTUALIZADO: {doc_id}")
    return {"ok": True}


@app.post("/api/documentos-base/{doc_id}/abrir")
def abrir_documento_base_endpoint(doc_id: str):
    from drive import (
        leer_registro_documentos_base,
        leer_datos_documento_base,
        leer_config_documento_base,
    )

    service = get_service()
    registro = leer_registro_documentos_base(service)
    doc_meta = None
    for doc in registro.get("documentos", []):
        if doc["id"] == doc_id:
            doc_meta = doc
            break

    if not doc_meta:
        raise HTTPException(status_code=404, detail="Documento base no encontrado")

    datos_cache = leer_datos_documento_base(service, doc_id)

    if not datos_cache or not datos_cache.get("texto"):
        raise HTTPException(status_code=400, detail="El documento base no tiene datos. Suba el archivo .docx primero.")

    texto = datos_cache.get("texto", "")
    config = leer_config_documento_base(service, doc_id)

    tiene_config = config is not None and config.get("mapeo") and len(config.get("mapeo", {})) > 0
    print(f"[MEMORIA] Abriendo documento base: {doc_id}")
    print(f"[MEMORIA] Configuración encontrada en Drive: {'SÍ' if tiene_config else 'NO'}")
    if tiene_config:
        print(f"[MEMORIA] Entradas en mapeo: {len(config.get('mapeo', {}))}")
    else:
        print(f"[MEMORIA] Se ejecutará detección dinámica de marcadores")

    caso_id = f"caso_{generar_hash(doc_id + datetime.datetime.now().isoformat())[:8]}"
    caso_documento_id = f"caso_{generar_hash(caso_id)}"

    return {
        "mensaje": "Documento base cargado",
        "documento_id": caso_documento_id,
        "documento_base_id": doc_id,
        "caso_id": caso_id,
        "texto": texto,
        "texto_original": datos_cache.get("texto_original", texto),
        "config": config,
        "nombre_documento_base": doc_meta.get("nombre", ""),
    }


@app.post("/api/documentos-base/{doc_id}/guardar-configuracion")
def guardar_configuracion_documento_base_endpoint(doc_id: str, data: dict = Body(...)):
    from drive import (
        leer_registro_documentos_base,
        guardar_config_documento_base,
    )

    service = get_service()
    registro = leer_registro_documentos_base(service)
    encontrado = False
    for doc in registro.get("documentos", []):
        if doc["id"] == doc_id:
            encontrado = True
            break

    if not encontrado:
        raise HTTPException(status_code=404, detail="Documento base no encontrado")

    mapeo = data.get("mapeo", {})
    resaltados = data.get("resaltados", {})
    customBlocks = data.get("customBlocks", [])
    deletedCustomBlocks = data.get("deletedCustomBlocks", [])
    blockEntries = data.get("blockEntries", {})
    blockOrder = data.get("blockOrder", None)
    blockNames = data.get("blockNames", None)
    memoriaDocData = data.get("memoriaDocData", None)
    dynamicEntries = data.get("dynamicEntries", {})
    config = {
        "mapeo": mapeo,
        "resaltados": resaltados,
        "customBlocks": customBlocks,
        "deletedCustomBlocks": deletedCustomBlocks,
        "blockEntries": blockEntries,
        "blockOrder": blockOrder,
        "blockNames": blockNames,
        "memoriaDocData": memoriaDocData,
        "dynamicEntries": dynamicEntries,
        "fecha_configuracion": datetime.datetime.now().isoformat(),
        "config_version": 2,
    }
    carpeta_doc = guardar_config_documento_base(service, doc_id, config)
    print(f"[MEMORIA] Guardando configuración: {doc_id} ({len(mapeo)} entradas, {len(resaltados)} resaltados)")
    print(f"[MEMORIA] Carpeta destino: {carpeta_doc}")
    print(f"[MEMORIA] Configuración guardada exitosamente en Google Drive")
    return {"ok": True, "carpeta_doc": carpeta_doc}


@app.post("/api/documentos-base/{doc_id}/reanalizar")
def reanalizar_documento_base_endpoint(doc_id: str):
    from drive import (
        leer_registro_documentos_base,
        leer_datos_documento_base,
        guardar_config_documento_base,
    )
    from ia import detectar_variables

    service = get_service()
    registro = leer_registro_documentos_base(service)
    doc_meta = None
    for doc in registro.get("documentos", []):
        if doc["id"] == doc_id:
            doc_meta = doc
            break

    if not doc_meta:
        raise HTTPException(status_code=404, detail="Documento base no encontrado")

    datos = leer_datos_documento_base(service, doc_id)
    texto = datos.get("texto", "") if datos else ""

    if not texto:
        raise HTTPException(status_code=400, detail="El documento base no tiene texto")

    try:
        variables = detectar_variables(texto)
    except Exception as e:
        print(f"⚠️ Error reanalizando: {e}")
        variables = []

    existing_config = leer_config_documento_base(service, doc_id) or {}
    config = dict(existing_config)
    config["variables"] = variables if variables else []
    config["fecha_configuracion"] = datetime.datetime.now().isoformat()
    config["config_version"] = max(config.get("config_version", 1), 2)
    guardar_config_documento_base(service, doc_id, config)
    print(f"🔄 REANALIZADO: {doc_id} ({len(variables)} variables)")
    return {"ok": True, "config": config}


@app.post("/api/documentos-base/{doc_id}/guardar-cambios")
async def guardar_cambios_documento_base_endpoint(doc_id: str, data: dict = Body(...)):
    from drive import (
        leer_registro_documentos_base,
        guardar_registro_documentos_base,
        subir_documento_base,
        guardar_datos_documento_base,
        guardar_config_documento_base,
        leer_config_documento_base,
        leer_datos_documento_base,
        obtener_carpeta_archivos_base,
        obtener_o_crear_carpeta,
        subir_o_actualizar_json,
    )
    from ia import detectar_variables

    texto_plano = data.get("texto_plano", "")
    if not texto_plano.strip():
        raise HTTPException(status_code=400, detail="No se proporcionó texto para guardar")

    lineas_tp = texto_plano.split("\n")
    while lineas_tp and lineas_tp[-1].strip() == "":
        lineas_tp.pop()
    texto_plano = "\n".join(lineas_tp)

    service = get_service()
    registro = leer_registro_documentos_base(service)
    documentos = registro.get("documentos", [])

    doc_meta = None
    for doc in documentos:
        if doc["id"] == doc_id:
            doc_meta = doc
            break

    if not doc_meta:
        raise HTTPException(status_code=404, detail="Documento base no encontrado")

    try:
        docx_doc = Document()
        for linea in texto_plano.split("\n"):
            docx_doc.add_paragraph(linea)
        import io
        buf = io.BytesIO()
        docx_doc.save(buf)
        file_bytes = buf.getvalue()
        if len(file_bytes) < 100:
            raise Exception("DOCX generado demasiado pequeño")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando DOCX: {str(e)}")

    current_version = doc_meta.get("version", 1)
    try:
        carpeta_base = obtener_carpeta_archivos_base(service)
        carpeta_doc = obtener_o_crear_carpeta(service, doc_id, carpeta_base)
        existing_datos = leer_datos_documento_base(service, doc_id)
        if existing_datos:
            version_filename = f"datos.json_v{current_version}"
            subir_o_actualizar_json(service, carpeta_doc, version_filename, existing_datos)
            print(f"📦 Versión anterior guardada: {version_filename}")
    except Exception as e:
        print(f"⚠️ Error guardando versión anterior: {e}")

    archivo_nombre = doc_meta.get("archivo_nombre", f"documento_{doc_id}.docx")
    subir_documento_base(service, doc_id, file_bytes, archivo_nombre)

    guardar_datos_documento_base(service, doc_id, {
        "texto": texto_plano,
        "texto_original": texto_plano,
        "prompt_usado": "",
        "prompt_version": PROMPT_VERSION,
    })

    try:
        variables = detectar_variables(texto_plano)
        existing_config = leer_config_documento_base(service, doc_id) or {}
        config = dict(existing_config)
        config["variables"] = variables if variables else []
        config["fecha_configuracion"] = datetime.datetime.now().isoformat()
        config["config_version"] = max(config.get("config_version", 1), 2)
        guardar_config_documento_base(service, doc_id, config)
    except Exception as e:
        print(f"⚠️ Error re-detectando variables: {e}")

    new_version = current_version + 1
    for doc in documentos:
        if doc["id"] == doc_id:
            doc["version"] = new_version
            doc["fecha_modificacion"] = datetime.datetime.now().isoformat()
            break
    guardar_registro_documentos_base(service, registro)

    print(f"✅ CAMBIOS GUARDADOS: {doc_id} (v{current_version} → v{new_version})")
    return {"ok": True, "version": new_version, "anterior_version": current_version}


def _agregarATaxonomia(registro, categoria, subcategoria=None, procedimiento=None):
    if not categoria:
        return
    taxonomy = registro.setdefault("taxonomy", {})
    cats_extra = taxonomy.setdefault("categorias_extra", {})
    cat_upper = categoria.upper().strip()
    if cat_upper not in cats_extra:
        cats_extra[cat_upper] = {}
    if subcategoria:
        sub_upper = subcategoria.strip()
        if sub_upper not in cats_extra[cat_upper]:
            cats_extra[cat_upper][sub_upper] = []
        if procedimiento:
            proc_upper = procedimiento.strip().upper()
            if proc_upper and proc_upper not in cats_extra[cat_upper][sub_upper]:
                cats_extra[cat_upper][sub_upper].append(proc_upper)
    procs_extra = taxonomy.setdefault("procedimientos_extra", [])
    if procedimiento:
        proc_upper = procedimiento.strip().upper()
        if proc_upper and proc_upper not in procs_extra:
            procs_extra.append(proc_upper)


@app.get("/api/documentos-base/taxonomy")
def obtener_taxonomia():
    from drive import leer_registro_documentos_base
    service = get_service()
    registro = leer_registro_documentos_base(service)
    taxonomy = registro.get("taxonomy", {"categorias_extra": {}, "procedimientos_extra": []})
    return taxonomy


@app.post("/api/documentos-base/taxonomy/add")
def agregar_taxonomia_endpoint(data: dict = Body(...)):
    from drive import (
        leer_registro_documentos_base,
        guardar_registro_documentos_base,
    )
    categoria = (data.get("categoria") or "").strip()
    subcategoria = (data.get("subcategoria") or "").strip() or None
    procedimiento = (data.get("procedimiento") or "").strip() or None
    if not categoria:
        raise HTTPException(status_code=400, detail="Categoría requerida")
    service = get_service()
    registro = leer_registro_documentos_base(service)
    _agregarATaxonomia(registro, categoria, subcategoria, procedimiento)
    guardar_registro_documentos_base(service, registro)
    return {"ok": True, "taxonomy": registro.get("taxonomy", {})}


@app.post("/api/documentos-base/seed")
def seed_documento_base():
    from drive import (
        leer_registro_documentos_base,
        guardar_registro_documentos_base,
    )

    service = get_service()
    registro = leer_registro_documentos_base(service)

    seed_id = generar_hash("demanda_servidumbre_paso_jurisflow")

    seeds_eliminados = registro.get("seeds_eliminados", [])
    if seed_id in seeds_eliminados:
        return {"ok": True, "mensaje": "Seed eliminado previamente"}

    for doc in registro.get("documentos", []):
        if doc["id"] == seed_id:
            return {"ok": True, "mensaje": "Ya existe"}

    seed_doc = {
        "id": seed_id,
        "nombre": "Demanda de Servidumbre de Paso",
        "categoria": "CIVIL",
        "subcategoria": "Servidumbre de paso",
        "procedimiento": "SUMARIO",
        "descripcion": "Documento base para demanda de constitución de servidumbre de paso.",
        "archivo_nombre": "Formato_Unico_Demanda_Servidumbre_de_Paso_JurisFlow.docx",
        "fecha_creacion": datetime.datetime.now().isoformat(),
        "fecha_modificacion": datetime.datetime.now().isoformat(),
        "version": 1,
    }

    registro["documentos"].append(seed_doc)
    guardar_registro_documentos_base(service, registro)

    print("📚 SEED: Demanda de Servidumbre de Paso registrado")
    return {"ok": True, "documento": seed_doc}